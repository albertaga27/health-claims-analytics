"""
Document Extraction Agent for Claims Health Analytics.
Uses Azure Document Intelligence and GPT models for comprehensive document processing.
"""
import os
import asyncio
import json
from typing import Any, Dict, List, Optional, Union
import io
from datetime import datetime
from PIL import Image
import pypdf
import docx
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import AnalyzeDocumentRequest
from azure.ai.inference import ChatCompletionsClient
from azure.ai.inference.models import SystemMessage, UserMessage
from azure.core.credentials import AzureKeyCredential

from base_agent import BaseAgent
from models import (
    DocumentMetadata, DocumentExtractionResult, ExtractedEntity
)
from utils import (
    timing_decorator, sanitize_text, generate_document_id,
    DocumentProcessingException, validate_file_type
)
from config import AzureConfig, AgentConfig

class DocumentExtractionAgent(BaseAgent):
    """Agent responsible for extracting and analyzing documents."""
    
    SUPPORTED_FILE_TYPES = ['pdf', 'docx', 'doc', 'txt', 'png', 'jpg', 'jpeg', 'tiff', 'bmp']
    
    def __init__(self, azure_config: AzureConfig, agent_config: AgentConfig):
        """Initialize the Document Extraction Agent."""
        super().__init__(azure_config, agent_config)
        self.chat_client = self._setup_chat_client()
        self.document_intelligence_client = self._setup_document_intelligence_client()
    
    def _setup_chat_client(self) -> ChatCompletionsClient:
        """Set up the chat client for GPT processing."""
        try:
            return ChatCompletionsClient(
                endpoint=self.azure_config.openai_endpoint,
                credential=self.credential
            )
        except Exception as e:
            self.logger.error(f"Failed to setup chat client: {e}")
            raise DocumentProcessingException(f"Chat client setup failed: {e}")
    
    def _setup_document_intelligence_client(self) -> Optional[DocumentIntelligenceClient]:
        """Set up the Document Intelligence client."""
        try:
            if (self.azure_config.document_intelligence_endpoint and 
                self.azure_config.document_intelligence_key):
                return DocumentIntelligenceClient(
                    endpoint=self.azure_config.document_intelligence_endpoint,
                    credential=AzureKeyCredential(self.azure_config.document_intelligence_key)
                )
            else:
                self.logger.warning("Document Intelligence not configured, using GPT-only processing")
                return None
        except Exception as e:
            self.logger.warning(f"Document Intelligence setup failed: {e}, falling back to GPT-only")
            return None
    
    @timing_decorator
    async def process(self, input_data: Union[DocumentMetadata, str, bytes]) -> DocumentExtractionResult:
        """Process a document and extract relevant information."""
        try:
            if isinstance(input_data, DocumentMetadata):
                return await self._process_document_metadata(input_data)
            elif isinstance(input_data, str):
                return await self._process_file_path(input_data)
            elif isinstance(input_data, bytes):
                return await self._process_document_bytes(input_data)
            else:
                raise DocumentProcessingException(f"Unsupported input type: {type(input_data)}")
                
        except Exception as e:
            self.logger.error(f"Document processing failed: {e}")
            raise DocumentProcessingException(f"Document processing failed: {e}")
    
    async def _process_document_metadata(self, metadata: DocumentMetadata) -> DocumentExtractionResult:
        """Process document using metadata information."""
        # For demo purposes, we'll simulate processing
        # In a real implementation, you would load the actual file
        
        extracted_text = f"Processed document: {metadata.file_name}"
        entities = [
            ExtractedEntity(
                text="Sample medical condition",
                category="HEALTH_CONDITION",
                confidence_score=0.85,
                offset=0,
                length=23
            )
        ]
        
        return DocumentExtractionResult(
            document_metadata=metadata,
            extracted_text=extracted_text,
            extracted_entities=entities,
            key_value_pairs={"document_type": "medical_record"},
            tables=[],
            confidence_score=0.85,
            processing_time=0.0  # Will be set by timing decorator
        )
    
    async def _process_file_path(self, file_path: str) -> DocumentExtractionResult:
        """Process document from file path."""
        if not os.path.exists(file_path):
            raise DocumentProcessingException(f"File not found: {file_path}")
        
        # Validate file type
        if not validate_file_type(file_path, self.SUPPORTED_FILE_TYPES):
            raise DocumentProcessingException(f"Unsupported file type: {file_path}")
        
        # Create metadata
        file_stats = os.stat(file_path)
        metadata = DocumentMetadata(
            file_name=os.path.basename(file_path),
            file_type=file_path.split('.')[-1].lower(),
            file_size=file_stats.st_size,
            upload_time=datetime.fromtimestamp(file_stats.st_ctime),
            document_id=generate_document_id()
        )
        
        # Extract text and structure based on file type
        extraction_result = {}
        if metadata.file_type == 'pdf':
            extraction_result = await self._extract_from_pdf(file_path)
        elif metadata.file_type in ['docx', 'doc']:
            extraction_result = await self._extract_from_docx(file_path)
        elif metadata.file_type == 'txt':
            extraction_result = await self._extract_from_txt(file_path)
        elif metadata.file_type in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
            extraction_result = await self._extract_from_image(file_path)
        else:
            raise DocumentProcessingException(f"Unsupported file type: {metadata.file_type}")
        
        # Update metadata with page count if available
        if 'page_count' in extraction_result:
            metadata.pages = extraction_result['page_count']
        
        # Process the extracted data with AI models (if needed for entity extraction)
        return await self._process_extraction_result(metadata, extraction_result)
    
    async def _process_document_bytes(self, document_bytes: bytes) -> DocumentExtractionResult:
        """Process document from bytes."""
        # Extract text and structure from bytes (simplified approach)
        extracted_text = document_bytes.decode('utf-8', errors='ignore')
        
        metadata = DocumentMetadata(
            file_name="document_from_bytes",
            file_type="unknown",
            file_size=len(document_bytes),
            upload_time=datetime.now(),
            document_id=generate_document_id()
        )
        
        extraction_result = {
            "text": sanitize_text(extracted_text),
            "tables": [],
            "key_value_pairs": {},
            "paragraphs": [{"content": line.strip(), "role": None, "bounding_regions": []} for line in extracted_text.split('\n') if line.strip()],
            "page_count": 1
        }
        
        return await self._process_extraction_result(metadata, extraction_result)
    
    async def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text, tables, and structure from PDF file using Document Intelligence Layout."""
        if self.document_intelligence_client:
            # Use Document Intelligence for better PDF processing with layout analysis
            return await self._extract_with_document_intelligence(file_path)
        else:
            # Fallback to pypdf for basic text extraction
            try:
                text = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                
                return {
                    "text": sanitize_text(text),
                    "tables": [],
                    "key_value_pairs": {},
                    "paragraphs": [],
                    "page_count": len(pdf_reader.pages)
                }
            except Exception as e:
                raise DocumentProcessingException(f"PDF extraction failed: {e}")
    
    async def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract tables from DOCX
            tables = []
            for table in doc.tables:
                table_data = {
                    "row_count": len(table.rows),
                    "column_count": len(table.columns) if table.rows else 0,
                    "cells": []
                }
                
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_data = {
                            "row_index": row_idx,
                            "column_index": col_idx,
                            "content": cell.text.strip(),
                            "row_span": 1,
                            "column_span": 1
                        }
                        table_data["cells"].append(cell_data)
                
                tables.append(table_data)
            
            return {
                "text": sanitize_text(text),
                "tables": tables,
                "key_value_pairs": {},
                "paragraphs": [{"content": para.text, "role": None, "bounding_regions": []} for para in doc.paragraphs if para.text.strip()],
                "page_count": 1
            }
        except Exception as e:
            raise DocumentProcessingException(f"DOCX extraction failed: {e}")

    async def _extract_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            
            return {
                "text": sanitize_text(text),
                "tables": [],
                "key_value_pairs": {},
                "paragraphs": [{"content": line.strip(), "role": None, "bounding_regions": []} for line in text.split('\n') if line.strip()],
                "page_count": 1
            }
        except Exception as e:
            raise DocumentProcessingException(f"TXT extraction failed: {e}")

    async def _extract_from_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text, tables, and structure from image using Document Intelligence Layout."""
        if self.document_intelligence_client:
            return await self._extract_with_document_intelligence(file_path)
        else:
            # Fallback to basic text extraction if Document Intelligence is not available
            return {
                "text": "Document Intelligence not configured - image processing unavailable",
                "tables": [],
                "key_value_pairs": {},
                "paragraphs": [],
                "page_count": 1
            }
    
    async def _extract_with_document_intelligence(self, file_path: str) -> Dict[str, Any]:
        """Extract text, tables, and key-value pairs using Azure Document Intelligence Layout model."""
        try:
            with open(file_path, 'rb') as file:
                # Use the Layout model for comprehensive document analysis
                poller = self.document_intelligence_client.begin_analyze_document(
                    "prebuilt-layout", 
                    analyze_request=file,
                    content_type="application/octet-stream"
                )
                result = poller.result()
                
                # Extract text content
                extracted_text = ""
                if result.content:
                    extracted_text = result.content
                
                # Extract tables
                tables = []
                if result.tables:
                    for table in result.tables:
                        table_data = {
                            "row_count": table.row_count,
                            "column_count": table.column_count,
                            "cells": []
                        }
                        
                        if table.cells:
                            for cell in table.cells:
                                cell_data = {
                                    "row_index": cell.row_index,
                                    "column_index": cell.column_index,
                                    "content": cell.content if cell.content else "",
                                    "row_span": getattr(cell, 'row_span', 1),
                                    "column_span": getattr(cell, 'column_span', 1)
                                }
                                table_data["cells"].append(cell_data)
                        
                        tables.append(table_data)
                
                # Extract key-value pairs
                key_value_pairs = {}
                if result.key_value_pairs:
                    for kv_pair in result.key_value_pairs:
                        if kv_pair.key and kv_pair.value:
                            key_text = kv_pair.key.content if kv_pair.key.content else ""
                            value_text = kv_pair.value.content if kv_pair.value.content else ""
                            if key_text and value_text:
                                key_value_pairs[key_text.strip()] = value_text.strip()
                
                # Extract paragraphs and their bounding information
                paragraphs = []
                if result.paragraphs:
                    for para in result.paragraphs:
                        paragraph_data = {
                            "content": para.content if para.content else "",
                            "role": getattr(para, 'role', None),
                            "bounding_regions": []
                        }
                        
                        if hasattr(para, 'bounding_regions') and para.bounding_regions:
                            for region in para.bounding_regions:
                                region_data = {
                                    "page_number": region.page_number if hasattr(region, 'page_number') else 1,
                                    "polygon": region.polygon if hasattr(region, 'polygon') else []
                                }
                                paragraph_data["bounding_regions"].append(region_data)
                        
                        paragraphs.append(paragraph_data)
                
                return {
                    "text": sanitize_text(extracted_text),
                    "tables": tables,
                    "key_value_pairs": key_value_pairs,
                    "paragraphs": paragraphs,
                    "page_count": len(result.pages) if result.pages else 1
                }
                
        except Exception as e:
            self.logger.error(f"Document Intelligence extraction failed: {e}")
            raise DocumentProcessingException(f"Document Intelligence extraction failed: {e}")
    
    async def _process_extraction_result(self, metadata: DocumentMetadata, extraction_result: Dict[str, Any]) -> DocumentExtractionResult:
        """Process the extraction result and create entities using AI models if needed."""
        try:
            extracted_text = extraction_result.get("text", "")
            tables = extraction_result.get("tables", [])
            key_value_pairs = extraction_result.get("key_value_pairs", {})
            paragraphs = extraction_result.get("paragraphs", [])
            
            # Extract entities from text using AI models (GPT) for medical entities
            entities = []
            if extracted_text.strip():
                entities = await self._extract_medical_entities_with_ai(extracted_text)
            
            # Add document type classification based on extracted content
            document_type = await self._classify_document_type(extracted_text, key_value_pairs)
            key_value_pairs["document_type"] = document_type
            
            # Calculate confidence score based on extraction quality
            confidence_score = await self._calculate_confidence_score(extraction_result)
            
            return DocumentExtractionResult(
                document_metadata=metadata,
                extracted_text=extracted_text,
                extracted_entities=entities,
                key_value_pairs=key_value_pairs,
                tables=tables,
                confidence_score=confidence_score,
                processing_time=0.0  # Set by timing decorator
            )
            
        except Exception as e:
            raise DocumentProcessingException(f"Extraction result processing failed: {e}")

    async def _extract_medical_entities_with_ai(self, text: str) -> List[ExtractedEntity]:
        """Extract medical entities from text using GPT for healthcare-specific analysis."""
        try:
            system_prompt = """You are a medical document analysis expert specializing in healthcare insurance claims.
Extract key medical entities from the provided text and return them in a structured format.
Focus on:
- Medical conditions and diagnoses
- Medications and treatments
- Medical procedures
- Symptoms and signs
- Body parts and anatomy
- Dates related to medical events
- Healthcare provider information

For each entity, provide the exact text, category, and confidence level."""

            user_prompt = f"""Extract medical entities from this text:

{text[:3000]}  # Limit to avoid token limits

Return entities in this JSON format:
[{{"text": "entity_text", "category": "MEDICAL_CONDITION", "confidence": 0.85, "offset": 0, "length": 11}}]"""

            response = await self.chat_client.complete(
                messages=[
                    SystemMessage(content=system_prompt),
                    UserMessage(content=user_prompt)
                ],
                model=self.agent_config.document_extraction_model,
                temperature=0.1,  # Lower temperature for more consistent extraction
                max_tokens=1000
            )
            
            response_text = response.choices[0].message.content
            
            # Parse the GPT response to extract entities
            return await self._parse_ai_entities_response(response_text, text)
            
        except Exception as e:
            self.logger.warning(f"AI entity extraction failed: {e}")
            # Return basic entities as fallback
            return await self._extract_basic_entities(text)

    async def _classify_document_type(self, text: str, key_value_pairs: Dict[str, str]) -> str:
        """Classify the document type based on content."""
        text_lower = text.lower()
        
        # Check for medical document indicators
        if any(keyword in text_lower for keyword in ['claim', 'insurance', 'policy', 'coverage']):
            return "insurance_claim"
        elif any(keyword in text_lower for keyword in ['diagnosis', 'treatment', 'medical', 'patient', 'doctor']):
            return "medical_record"
        elif any(keyword in text_lower for keyword in ['prescription', 'medication', 'dosage', 'pharmacy']):
            return "prescription"
        elif any(keyword in text_lower for keyword in ['lab', 'test', 'result', 'blood', 'urine']):
            return "lab_report"
        else:
            return "general_medical"

    async def _calculate_confidence_score(self, extraction_result: Dict[str, Any]) -> float:
        """Calculate confidence score based on extraction quality."""
        base_score = 0.5
        
        # Increase confidence if we have structured data
        if extraction_result.get("key_value_pairs"):
            base_score += 0.2
        
        if extraction_result.get("tables"):
            base_score += 0.15
            
        # Check text quality
        text = extraction_result.get("text", "")
        if len(text) > 100:
            base_score += 0.1
            
        # Cap at 0.95 to account for potential extraction errors
        return min(base_score, 0.95)

    async def _parse_ai_entities_response(self, response_text: str, original_text: str) -> List[ExtractedEntity]:
        """Parse the AI response to extract entities."""
        entities = []
        
        try:
            # Try to parse as JSON first
            if response_text.strip().startswith('['):
                entity_data = json.loads(response_text)
                for entity in entity_data:
                    entities.append(ExtractedEntity(
                        text=entity.get("text", ""),
                        category=entity.get("category", "UNKNOWN"),
                        confidence_score=float(entity.get("confidence", 0.5)),
                        offset=int(entity.get("offset", 0)),
                        length=int(entity.get("length", len(entity.get("text", ""))))
                    ))
        except (json.JSONDecodeError, ValueError, KeyError) as e:
            self.logger.warning(f"Failed to parse AI entities response as JSON: {e}")
            # Fallback to basic extraction
            entities = await self._extract_basic_entities(original_text)
        
        return entities

    async def _extract_basic_entities(self, text: str) -> List[ExtractedEntity]:
        """Extract basic medical entities using keyword matching as fallback."""
        entities = []
        text_lower = text.lower()
        
        # Basic medical terms to look for
        medical_terms = {
            'diabetes': 'MEDICAL_CONDITION',
            'hypertension': 'MEDICAL_CONDITION',
            'cancer': 'MEDICAL_CONDITION',
            'surgery': 'PROCEDURE',
            'medication': 'MEDICATION',
            'treatment': 'PROCEDURE',
            'diagnosis': 'MEDICAL_CONDITION',
            'prescription': 'MEDICATION',
            'doctor': 'HEALTHCARE_PROVIDER',
            'hospital': 'HEALTHCARE_FACILITY'
        }
        
        for term, category in medical_terms.items():
            offset = text_lower.find(term)
            if offset >= 0:
                entities.append(ExtractedEntity(
                    text=term,
                    category=category,
                    confidence_score=0.7,
                    offset=offset,
                    length=len(term)
                ))
        
        return entities

    async def process_multiple_documents(self, documents: List[str]) -> List[DocumentExtractionResult]:
        """Process multiple documents concurrently."""
        tasks = [self.process(doc_path) for doc_path in documents]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                self.logger.error(f"Failed to process document {documents[i]}: {result}")
                # Create error result
                error_result = DocumentExtractionResult(
                    document_metadata=DocumentMetadata(
                        file_name=os.path.basename(documents[i]),
                        file_type="unknown",
                        file_size=0,
                        upload_time=datetime.now(),
                        document_id=generate_document_id()
                    ),
                    extracted_text="",
                    extracted_entities=[],
                    key_value_pairs={},
                    tables=[],
                    confidence_score=0.0,
                    processing_time=0.0,
                    errors=[str(result)]
                )
                processed_results.append(error_result)
            else:
                processed_results.append(result)
        
        return processed_results